"""Extracción de documentos a HTML estructurado (PDF, Word, texto).

Usa PyMuPDF (fitz), ya presente en el backend, para leer PDFs.
"""
import io
import logging
from pathlib import Path

from .html_formatter import asegurar_html, envolver_documento, estructurar_texto_legal

logger = logging.getLogger(__name__)

EXTENSIONES_PERMITIDAS = {
    '.txt',
    '.md',
    '.html',
    '.htm',
    '.pdf',
    '.docx',
}


def extension_permitida(nombre: str) -> bool:
    return Path(nombre).suffix.lower() in EXTENSIONES_PERMITIDAS


def extraer_texto_archivo(archivo) -> tuple[str, str]:
    """Extrae contenido como HTML estructurado. Retorna (html, extensión)."""
    nombre = getattr(archivo, 'name', 'documento.txt')
    ext = Path(nombre).suffix.lower()

    if hasattr(archivo, 'read'):
        raw = archivo.read()
        if hasattr(archivo, 'seek'):
            archivo.seek(0)
    else:
        raw = archivo

    if ext == '.pdf':
        return _extraer_pdf_html(raw), ext
    if ext == '.docx':
        return _extraer_docx_html(raw), ext
    if ext in ('.html', '.htm'):
        return asegurar_html(_leer_texto(raw)), ext
    return asegurar_html(_leer_texto(raw)), ext


def _leer_texto(raw: bytes | str) -> str:
    if isinstance(raw, str):
        return raw
    for enc in ('utf-8', 'latin-1', 'cp1252'):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode('utf-8', errors='replace')


def _extraer_pdf_html(raw: bytes) -> str:
    import fitz  # PyMuPDF

    paginas: list[str] = []
    with fitz.open(stream=raw, filetype='pdf') as doc:
        for page in doc:
            texto = page.get_text('text')
            if texto and texto.strip():
                paginas.append(texto)
    if not paginas:
        logger.warning('PDF sin texto extraíble')
        return asegurar_html('')
    plano = '\n\n'.join(paginas)
    return estructurar_texto_legal(plano)


def _runs_a_html(paragraph) -> str:
    """Mapea runs de Word a negritas."""
    partes = []
    for run in paragraph.runs:
        t = run.text
        if not t:
            continue
        if run.bold:
            partes.append(f'<strong>{html_escape(t)}</strong>')
        else:
            partes.append(html_escape(t))
    return ''.join(partes) if partes else html_escape(paragraph.text)


def html_escape(t: str) -> str:
    import html as html_mod

    return html_mod.escape(t, quote=False)


def _estilo_parrafo_docx(paragraph) -> str:
    style_name = (paragraph.style.name or '').lower() if paragraph.style else ''
    texto = _runs_a_html(paragraph).strip()
    if not texto:
        return ''

    if 'heading 1' in style_name or 'título 1' in style_name:
        return f'<h1 class="ley-titulo">{texto}</h1>'
    if 'heading 2' in style_name or 'título 2' in style_name:
        return f'<h2 class="ley-capitulo">{texto}</h2>'
    if 'heading 3' in style_name or 'título 3' in style_name:
        return f'<h3 class="ley-articulo">{texto}</h3>'
    if 'heading' in style_name or 'título' in style_name:
        return f'<h4 class="ley-subtitulo">{texto}</h4>'

    if paragraph._p.pPr is not None:
        numPr = paragraph._p.pPr.numPr
        if numPr is not None:
            return f'<li class="ley-lista">{texto}</li>'

    return f'<p class="ley-parrafo">{texto}</p>'


def _extraer_docx_html(raw: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(raw))
    bloques: list[str] = []
    en_lista = False

    for paragraph in doc.paragraphs:
        texto_plano = paragraph.text.strip()
        if not texto_plano:
            if en_lista:
                bloques.append('</ul>')
                en_lista = False
            continue

        style_name = (paragraph.style.name or '').lower() if paragraph.style else ''
        es_lista = 'list' in style_name or (
            paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
        )

        if es_lista:
            if not en_lista:
                bloques.append('<ul class="ley-lista-ul">')
                en_lista = True
            bloques.append(f'<li class="ley-lista">{_runs_a_html(paragraph)}</li>')
        else:
            if en_lista:
                bloques.append('</ul>')
                en_lista = False
            bloque = _estilo_parrafo_docx(paragraph)
            if not bloque:
                bloque = f'<p class="ley-parrafo">{_runs_a_html(paragraph)}</p>'
            bloques.append(bloque)

    if en_lista:
        bloques.append('</ul>')

    for tabla in doc.tables:
        bloques.append('<table class="ley-tabla"><tbody>')
        for fila in tabla.rows:
            bloques.append('<tr>')
            for celda in fila.cells:
                celdas_html = [
                    f'<p class="ley-parrafo">{c.text.strip()}</p>'
                    for c in celda.paragraphs
                    if c.text.strip()
                ]
                contenido_celda = ''.join(celdas_html) or '&nbsp;'
                bloques.append(f'<td>{contenido_celda}</td>')
            bloques.append('</tr>')
        bloques.append('</tbody></table>')

    cuerpo = '\n'.join(bloques) if bloques else '<p class="ley-parrafo"></p>'
    return envolver_documento(cuerpo)
