import hashlib
import re
import shutil
import tempfile
import time
import zipfile
from io import BytesIO
from pathlib import Path

from django.conf import settings
from django.core.files import File
from django.db import transaction
from django.utils import timezone

from .models import (
    AlertaCalidad,
    ArchivoDocumento,
    Documento,
    EvaluacionCalidad,
    HistorialDocumento,
    ResultadoConversion,
)
from .pdf_analysis import pdf_tiene_paginas_dominadas_por_imagen
from .services import registrar_historial


class ErrorConversion(Exception):
    def __init__(self, codigo, mensaje, *, detalles=None):
        super().__init__(mensaje)
        self.codigo = codigo
        self.mensaje = mensaje
        self.detalles = detalles or {}


def _limpiar_componente(valor):
    texto = re.sub(r'[<>:"/\\|?*\x00-\x1f]', ' ', str(valor or ''))
    texto = texto.replace(';', ',')
    return re.sub(r'\s+', ' ', texto).strip(' .')


def _componente_recortado(valor, limite):
    if len(valor) <= limite:
        return valor
    if limite <= 3:
        return valor[:limite]
    return valor[:limite - 3].rstrip() + '...'


def _limite_nombre_archivo(carpeta):
    limite_ruta = getattr(settings, 'FINAL_PATH_MAX_LENGTH', 240)
    reserva_raiz = max(
        len(str(settings.FINAL_NORMATIVA_ROOT)),
        getattr(settings, 'FINAL_STORAGE_ROOT_LENGTH_BUDGET', 100),
    )
    presupuesto_ruta = limite_ruta - reserva_raiz - len(carpeta) - 2
    if presupuesto_ruta < 80:
        raise ErrorConversion(
            'RUTA_FINAL_DEMASIADO_LARGA',
            'La carpeta de materia no permite crear un nombre de archivo seguro.',
        )
    return min(
        getattr(settings, 'FINAL_FILENAME_MAX_LENGTH', 230),
        presupuesto_ruta,
    )


def generar_nomenclatura_final(documento):
    componentes = [
        documento.efecto_normativo.abreviatura_archivo,
        documento.tipo_norma.abreviatura_archivo,
        documento.numero,
        documento.fecha_emision.strftime('%d-%m-%Y'),
        documento.titulo,
        documento.objeto,
        documento.materia.nombre,
    ]
    limpios = [_limpiar_componente(item) for item in componentes]
    nomenclatura = '; '.join(limpios)
    carpeta = _limpiar_componente(documento.materia.carpeta_destino)
    maximo = _limite_nombre_archivo(carpeta)
    nombre = f'{nomenclatura}.docx'
    truncado = False
    if len(nombre) > maximo:
        truncado = True
        digest = hashlib.sha256(nomenclatura.encode('utf-8')).hexdigest()[:8]
        limpios[5] = _componente_recortado(limpios[5], 55)
        limpios[4] = _componente_recortado(limpios[4], 65)
        limpios[6] = _componente_recortado(limpios[6], 40)
        nombre = f'{"; ".join(limpios)}; {digest}.docx'
        minimos = {5: 15, 4: 20, 6: 15, 2: 8}
        while len(nombre) > maximo:
            reducibles = [indice for indice, minimo in minimos.items() if len(limpios[indice]) > minimo]
            if not reducibles:
                break
            indice = max(reducibles, key=lambda item: len(limpios[item]) - minimos[item])
            limpios[indice] = _componente_recortado(limpios[indice], len(limpios[indice]) - 1)
            nombre = f'{"; ".join(limpios)}; {digest}.docx'
        if len(nombre) > maximo:
            sufijo = f'; {digest}.docx'
            nombre = f'{nombre[:maximo - len(sufijo)].rstrip(" .;")}{sufijo}'
    return nomenclatura, nombre, truncado


def _validar_ficha(documento):
    faltantes = []
    for campo in ('tipo_norma', 'efecto_normativo', 'materia', 'entidad_emisora'):
        if not getattr(documento, f'{campo}_id'):
            faltantes.append(campo)
    if documento.tipo_norma_id and documento.tipo_norma.requiere_numero and not documento.numero.strip():
        faltantes.append('numero')
    if documento.tipo_norma_id and documento.tipo_norma.requiere_fecha and not documento.fecha_emision:
        faltantes.append('fecha_emision')
    if not documento.titulo.strip():
        faltantes.append('titulo')
    if not documento.objeto.strip():
        faltantes.append('objeto')
    if faltantes:
        raise ErrorConversion(
            'FICHA_INCOMPLETA',
            'La ficha jurídica definitiva no está completa.',
            detalles={'campos_faltantes': faltantes},
        )


def _convertir_pdf(ruta_pdf, ruta_docx):
    try:
        from pdf2docx import Converter

        convertidor = Converter(str(ruta_pdf))
        try:
            convertidor.convert(str(ruta_docx), start=0, end=None)
        finally:
            convertidor.close()
    except Exception as exc:
        raise ErrorConversion(
            'CONVERSION_FALLIDA',
            'No se pudo convertir el PDF procesado a Word.',
            detalles={'error_conversor': str(exc)[:2000]},
        ) from exc
    if not ruta_docx.exists() or ruta_docx.stat().st_size <= 0 or not zipfile.is_zipfile(ruta_docx):
        raise ErrorConversion(
            'DOCX_INVALIDO',
            'El conversor no generó un archivo Word válido.',
        )
    try:
        from docx import Document

        documento_word = Document(str(ruta_docx))
        if not documento_word.sections:
            raise ValueError('El Word no contiene secciones.')
    except Exception as exc:
        raise ErrorConversion(
            'DOCX_INVALIDO',
            'El archivo Word generado no puede abrirse correctamente.',
        ) from exc


def _convertir_pdf_como_imagenes(ruta_pdf, ruta_docx):
    """Render every PDF page into Word without altering its visual layout."""
    try:
        import fitz
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt

        dpi = getattr(settings, 'PDF_WORD_IMAGE_DPI', 300)
        documento_word = Document()
        seccion = documento_word.sections[0]
        margen = Mm(2)
        seccion.top_margin = margen
        seccion.bottom_margin = margen
        seccion.left_margin = margen
        seccion.right_margin = margen
        seccion.header_distance = Mm(0)
        seccion.footer_distance = Mm(0)

        with fitz.open(ruta_pdf) as pdf:
            if not pdf.page_count:
                raise ErrorConversion('PDF_SIN_PAGINAS', 'El PDF no contiene pÃ¡ginas.')
            numero_paginas = pdf.page_count

            primera = pdf[0]
            seccion.page_width = Pt(primera.rect.width)
            seccion.page_height = Pt(primera.rect.height)
            ancho_disponible = seccion.page_width - seccion.left_margin - seccion.right_margin
            alto_disponible = seccion.page_height - seccion.top_margin - seccion.bottom_margin - Pt(2)

            for indice, pagina in enumerate(pdf):
                pixmap = pagina.get_pixmap(dpi=dpi, colorspace=fitz.csRGB, alpha=False)
                imagen = BytesIO(pixmap.tobytes('png'))
                escala = min(
                    ancho_disponible / Pt(pagina.rect.width),
                    alto_disponible / Pt(pagina.rect.height),
                )
                parrafo = documento_word.add_paragraph()
                parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                parrafo.paragraph_format.space_before = Pt(0)
                parrafo.paragraph_format.space_after = Pt(0)
                parrafo.add_run().add_picture(
                    imagen,
                    width=Pt(pagina.rect.width * escala),
                    height=Pt(pagina.rect.height * escala),
                )
                if indice < pdf.page_count - 1:
                    documento_word.add_page_break()

        documento_word.core_properties.title = 'Documento jurÃ­dico fiel al PDF original'
        documento_word.core_properties.subject = 'PÃ¡ginas conservadas como imÃ¡genes de alta calidad'
        documento_word.save(str(ruta_docx))
    except ErrorConversion:
        raise
    except Exception as exc:
        raise ErrorConversion(
            'CONVERSION_VISUAL_FALLIDA',
            'No se pudo crear el Word visual a partir del PDF original.',
            detalles={'error_conversor': str(exc)[:2000]},
        ) from exc

    if not ruta_docx.exists() or ruta_docx.stat().st_size <= 0 or not zipfile.is_zipfile(ruta_docx):
        raise ErrorConversion('DOCX_INVALIDO', 'El generador visual no creÃ³ un Word vÃ¡lido.')
    return {
        'conversor': 'pdf-imagen-alta-calidad',
        'texto_editable': False,
        'paginas_origen': numero_paginas,
        'calidad_dpi': dpi,
    }


def _limpiar_texto_word(texto):
    """Remove characters that are invalid in WordprocessingML."""
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', texto or '')


def _convertir_texto_ocr(resultado_procesamiento, ruta_docx):
    """Build an editable DOCX from the text already recognized during OCR."""
    try:
        from docx import Document
        from docx.shared import Cm, Pt

        paginas = list(resultado_procesamiento.paginas.order_by('numero_pagina'))
        if not paginas or not any(_limpiar_texto_word(pagina.texto).strip() for pagina in paginas):
            raise ErrorConversion(
                'OCR_TEXTO_VACIO',
                'El OCR no produjo texto suficiente para crear un Word editable.',
            )

        documento_word = Document()
        for seccion in documento_word.sections:
            seccion.top_margin = Cm(2)
            seccion.bottom_margin = Cm(2)
            seccion.left_margin = Cm(2.5)
            seccion.right_margin = Cm(2.5)

        estilo = documento_word.styles['Normal']
        estilo.font.name = 'Arial'
        estilo.font.size = Pt(11)

        paginas_con_texto = 0
        caracteres = 0
        for indice, pagina in enumerate(paginas):
            texto = _limpiar_texto_word(pagina.texto).strip()
            if texto:
                paginas_con_texto += 1
                caracteres += len(texto)
                bloques = [bloque.strip() for bloque in re.split(r'\n\s*\n', texto) if bloque.strip()]
                for bloque in bloques:
                    parrafo = documento_word.add_paragraph()
                    lineas = [linea.rstrip() for linea in bloque.splitlines()]
                    for numero_linea, linea in enumerate(lineas):
                        if numero_linea:
                            parrafo.add_run().add_break()
                        parrafo.add_run(linea)
            else:
                documento_word.add_paragraph('[Página sin texto reconocido]')
            if indice < len(paginas) - 1:
                documento_word.add_page_break()

        documento_word.core_properties.title = 'Documento jurídico editable generado mediante OCR'
        documento_word.core_properties.subject = 'Texto editable reconstruido desde un PDF escaneado'
        documento_word.save(str(ruta_docx))
    except ErrorConversion:
        raise
    except Exception as exc:
        raise ErrorConversion(
            'CONVERSION_OCR_FALLIDA',
            'No se pudo crear el Word editable a partir del texto OCR.',
            detalles={'error_conversor': str(exc)[:2000]},
        ) from exc

    if not ruta_docx.exists() or ruta_docx.stat().st_size <= 0 or not zipfile.is_zipfile(ruta_docx):
        raise ErrorConversion(
            'DOCX_INVALIDO',
            'El generador OCR no creó un archivo Word válido.',
        )
    return {
        'conversor': 'ocr-texto-editable',
        'texto_editable': True,
        'paginas_origen': len(paginas),
        'paginas_con_texto': paginas_con_texto,
        'caracteres_editables': caracteres,
    }


def _hash_archivo(ruta):
    digest = hashlib.sha256()
    with ruta.open('rb') as archivo:
        for bloque in iter(lambda: archivo.read(1024 * 1024), b''):
            digest.update(bloque)
    return digest.hexdigest()


def _nombre_disponible(resultado, carpeta, nombre):
    storage = resultado._meta.get_field('archivo').storage
    candidato = nombre
    version = 1
    stem, extension = Path(nombre).stem, Path(nombre).suffix
    maximo = _limite_nombre_archivo(carpeta)
    while storage.exists(f'{carpeta}/{candidato}'):
        version += 1
        sufijo = f'; v{version}{extension}'
        candidato = f'{stem[:maximo - len(sufijo)].rstrip(" .;")}{sufijo}'
    return candidato, version


def _marcar_error(documento_id, error, inicio):
    with transaction.atomic():
        documento = Documento.objects.select_for_update().get(pk=documento_id)
        estado_anterior = documento.estado
        resultado, _ = ResultadoConversion.objects.get_or_create(documento=documento)
        resultado.estado = ResultadoConversion.Estado.ERROR
        resultado.finalizado_at = timezone.now()
        resultado.duracion_ms = int((time.monotonic() - inicio) * 1000)
        resultado.error_codigo = getattr(error, 'codigo', 'ERROR_INTERNO')
        resultado.error_mensaje = getattr(error, 'mensaje', 'Ocurrió un error interno durante la conversión.')
        resultado.detalles_tecnicos = getattr(error, 'detalles', {})
        resultado.save()
        documento.estado = (
            Documento.Estado.PENDIENTE_APROBACION
            if resultado.error_codigo == 'FICHA_INCOMPLETA'
            else Documento.Estado.ERROR_CONVERSION
        )
        documento.save(update_fields=('estado', 'updated_at'))
        registrar_historial(
            documento,
            HistorialDocumento.Accion.CONVERSION_ERROR,
            estado_anterior=estado_anterior,
            estado_nuevo=documento.estado,
            descripcion=f'{resultado.error_codigo}: {resultado.error_mensaje}',
        )


def marcar_error_conversion(documento_id, error):
    resultado = ResultadoConversion.objects.filter(documento_id=documento_id).first()
    if (
        resultado
        and resultado.estado == ResultadoConversion.Estado.ERROR
        and resultado.error_codigo == getattr(error, 'codigo', '')
    ):
        return resultado
    _marcar_error(documento_id, error, time.monotonic())
    return ResultadoConversion.objects.get(documento_id=documento_id)


def convertir_documento_final(documento_id):
    inicio = time.monotonic()
    with transaction.atomic():
        documento = Documento.objects.select_for_update().get(pk=documento_id)
        resultado, _ = ResultadoConversion.objects.get_or_create(documento=documento)
        if resultado.estado == ResultadoConversion.Estado.COMPLETADA and resultado.archivo:
            return resultado
        if documento.estado not in {
            Documento.Estado.LISTO_PARA_CONVERSION,
            Documento.Estado.ERROR_CONVERSION,
        }:
            raise ErrorConversion(
                'ESTADO_INVALIDO',
                'El documento no está habilitado para conversión final.',
            )
        if documento.documento_canonico_id:
            raise ErrorConversion(
                'DOCUMENTO_DUPLICADO',
                'Un documento duplicado no puede generar otro archivo final.',
            )
        try:
            evaluacion = documento.evaluacion_calidad
        except EvaluacionCalidad.DoesNotExist as exc:
            raise ErrorConversion('SIN_CONTROL_CALIDAD', 'El documento no tiene control de calidad.') from exc
        alertas_pendientes = evaluacion.alertas.filter(
            estado=AlertaCalidad.Estado.ACTIVA,
        ).exists()
        if alertas_pendientes:
            raise ErrorConversion(
                'CALIDAD_NO_APROBADA',
                'Existen alertas pendientes de revisión.',
            )
        _validar_ficha(documento)
        anterior = documento.estado
        documento.estado = Documento.Estado.CONVIRTIENDO
        documento.save(update_fields=('estado', 'updated_at'))
        resultado.estado = ResultadoConversion.Estado.CONVIRTIENDO
        resultado.intentos += 1
        resultado.iniciado_at = timezone.now()
        resultado.finalizado_at = None
        resultado.error_codigo = ''
        resultado.error_mensaje = ''
        resultado.detalles_tecnicos = {}
        resultado.save()
        registrar_historial(
            documento,
            HistorialDocumento.Accion.CONVERSION_INICIADA,
            estado_anterior=anterior,
            estado_nuevo=documento.estado,
            descripcion=f'Conversión final iniciada (intento {resultado.intentos}).',
        )

    try:
        procesado = documento.archivos.get(
            tipo_archivo=ArchivoDocumento.TipoArchivo.PDF_PROCESADO,
        )
        original = documento.archivos.get(
            tipo_archivo=ArchivoDocumento.TipoArchivo.PDF_ORIGINAL,
        )
        with tempfile.TemporaryDirectory(prefix='conversion-final-') as temporal:
            temporal = Path(temporal)
            entrada = temporal / 'procesado.pdf'
            salida = temporal / 'normativa.docx'
            # The final Word must preserve the source appearance, so render the
            # untouched original PDF rather than the OCR-processed derivative.
            with original.archivo.open('rb') as origen, entrada.open('wb') as destino:
                shutil.copyfileobj(origen, destino)
            procesamiento = documento.resultado_procesamiento
            # Older records may have been classified as native PDFs because a
            # scanned image already contained a selectable OCR text layer.
            # Avoid pdf2docx in that case: it embeds every full-page image and
            # produces a huge, blurrier Word file.
            respaldado_por_imagen = pdf_tiene_paginas_dominadas_por_imagen(entrada)
            if procesamiento.requirio_ocr or respaldado_por_imagen:
                detalles_conversion = _convertir_pdf_como_imagenes(entrada, salida)
                detalles_conversion['pdf_respaldado_por_imagen'] = respaldado_por_imagen
            else:
                _convertir_pdf(entrada, salida)
                detalles_conversion = {
                    'conversor': 'pdf2docx',
                    'texto_editable': True,
                    'paginas_origen': procesamiento.numero_paginas,
                }
            nomenclatura, nombre_base, truncado = generar_nomenclatura_final(documento)
            carpeta = _limpiar_componente(documento.materia.carpeta_destino)
            if not carpeta:
                raise ErrorConversion('CARPETA_INVALIDA', 'La materia no tiene una carpeta final válida.')

            archivo_guardado = None
            try:
                with transaction.atomic():
                    documento = Documento.objects.select_for_update().get(pk=documento_id)
                    resultado = ResultadoConversion.objects.select_for_update().get(documento=documento)
                    nombre, version = _nombre_disponible(resultado, carpeta, nombre_base)
                    with salida.open('rb') as contenido:
                        resultado.carpeta_materia = carpeta
                        resultado.archivo.save(nombre, File(contenido), save=False)
                    archivo_guardado = resultado.archivo
                    resultado.estado = ResultadoConversion.Estado.COMPLETADA
                    resultado.nomenclatura_completa = nomenclatura
                    resultado.nombre_archivo = nombre
                    resultado.ruta_relativa = resultado.archivo.name
                    resultado.hash_sha256 = _hash_archivo(salida)
                    resultado.tamano_bytes = salida.stat().st_size
                    resultado.version = version
                    resultado.finalizado_at = timezone.now()
                    resultado.duracion_ms = int((time.monotonic() - inicio) * 1000)
                    resultado.detalles_tecnicos = {
                        **detalles_conversion,
                        'nombre_truncado': truncado,
                        'pdf_procesado_hash': procesado.hash_sha256,
                    }
                    resultado.save()
                    documento.estado = Documento.Estado.FINALIZADO
                    documento.save(update_fields=('estado', 'updated_at'))
                    registrar_historial(
                        documento,
                        HistorialDocumento.Accion.CONVERSION_COMPLETADA,
                        estado_anterior=Documento.Estado.CONVIRTIENDO,
                        estado_nuevo=documento.estado,
                        descripcion=f'Word final guardado en {resultado.ruta_relativa}.',
                    )
                return resultado
            except Exception:
                if archivo_guardado:
                    archivo_guardado.delete(save=False)
                raise
    except Exception as exc:
        error = exc if isinstance(exc, ErrorConversion) else ErrorConversion(
            'ERROR_INTERNO',
            'Ocurrió un error interno durante la conversión final.',
        )
        _marcar_error(documento_id, error, inicio)
        raise error from exc
