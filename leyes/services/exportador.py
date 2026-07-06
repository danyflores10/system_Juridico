"""Generación de documentos descargables (DOCX y PDF)."""

import io
import re
from datetime import datetime

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from .html_formatter import html_a_texto_plano


def _html_a_parrafos_docx(contenido: str) -> list[tuple[str, str]]:
    """Convierte HTML a lista (estilo, texto) para Word."""
    soup = BeautifulSoup(contenido, "html.parser")
    items = []
    for tag in soup.find_all(["h1", "h2", "h3", "h4", "p", "li"]):
        t = tag.get_text().strip()
        if not t:
            continue
        if tag.name == "h1":
            items.append(("heading1", t))
        elif tag.name == "h2":
            items.append(("heading2", t))
        elif tag.name == "h3":
            items.append(("heading3", t))
        else:
            items.append(("normal", t))
    if not items:
        plano = html_a_texto_plano(contenido)
        for bloque in plano.split("\n\n"):
            if bloque.strip():
                items.append(("normal", bloque.strip()))
    return items


def generar_docx(contenido: str, titulo: str, codigo_ley: str) -> bytes:
    doc = Document()
    doc.add_heading(f"Ley {codigo_ley} — {titulo}"[:200], level=0)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')} · LEGISLACION PROYECTO"
    ).italic = True
    doc.add_paragraph()
    estilo = doc.styles["Normal"]
    estilo.font.size = Pt(11)
    estilo.font.name = "Times New Roman"

    for tipo, texto in _html_a_parrafos_docx(contenido):
        if tipo == "heading1":
            doc.add_heading(texto, level=1)
        elif tipo == "heading2":
            doc.add_heading(texto, level=2)
        elif tipo == "heading3":
            doc.add_heading(texto, level=3)
        else:
            doc.add_paragraph(texto)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generar_pdf(contenido: str, titulo: str, codigo_ley: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, title=f"Ley_{codigo_ley}")
    estilos = getSampleStyleSheet()
    story = [
        Paragraph(f"<b>Ley {codigo_ley}</b>", estilos["Title"]),
        Paragraph(titulo[:300], estilos["Heading2"]),
        Spacer(1, 12),
    ]
    soup = BeautifulSoup(contenido, "html.parser")
    tags = soup.find_all(["h1", "h2", "h3", "h4", "p", "li"])
    if tags:
        for tag in tags:
            t = tag.get_text().strip()
            if not t:
                continue
            safe = (
                t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            )
            if tag.name == "h1":
                story.append(Paragraph(f"<b>{safe}</b>", estilos["Heading1"]))
            elif tag.name in ("h2", "h3"):
                story.append(Paragraph(f"<b>{safe}</b>", estilos["Heading2"]))
            else:
                story.append(Paragraph(safe, estilos["Normal"]))
            story.append(Spacer(1, 4))
    else:
        for parrafo in html_a_texto_plano(contenido).split("\n"):
            if parrafo.strip():
                safe = (
                    parrafo.strip()
                    .replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                )
                story.append(Paragraph(safe, estilos["Normal"]))
    doc.build(story)
    return buf.getvalue()
